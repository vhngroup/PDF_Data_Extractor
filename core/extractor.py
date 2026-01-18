import os
import fitz
import pdfplumber
import pandas as pd
from pdf2docx import Converter
from docx import Document
import re
import pytesseract
from PIL import Image
import io

# --- IMPORTACIONES GMFT ---
from gmft.pdf_bindings import PyPDFium2Document
from gmft.auto import AutoTableFormatter

class PDFExtractor:
    def __init__(self, file_path, output_folder="output"):
        self.file_path = file_path
        self.filename = os.path.splitext(os.path.basename(file_path))[0]
        self.output_dir = os.path.join(output_folder, self.filename)
        os.makedirs(self.output_dir, exist_ok=True)
        
        # --- CONFIGURACIÓN TESSERACT ---
        self.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        
        if os.path.exists(self.tesseract_cmd):
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
            self.has_ocr = True
            print("Motor de Visión Tesseract listo.")
        else:
            self.has_ocr = False
            print(f"Tesseract no encontrado, esta instalado?")

        # Configuración GMFT
        try:
            self.formatter = AutoTableFormatter()
            self.usar_ia = True
        except:
            self.usar_ia = False

    def _clean_text(self, text):
        if text:
            text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text) 
            text = text.strip()
            return text if text else None
        return None

    def _has_text_content(self):
        """Verifica si el PDF tiene texto seleccionable (Digital)"""
        try:
            with fitz.open(self.file_path) as doc:
                # Revisamos hasta 3 páginas
                for i in range(min(len(doc), 3)):
                    text = doc[i].get_text()
                    if len(text.strip()) > 50:
                        return True
            return False
        except:
            return False

    def extract_tables(self):
        """
        Lógica Inteligente Corregida:
        - Si es Digital -> Prioridad: Nativo > IA
        - Si es Imagen  -> Prioridad: IA > OCR
        """
        excel_path = os.path.join(self.output_dir, f"{self.filename}_tablas.xlsx")
        tables_found = []
        
        es_digital = self._has_text_content()

        # --- CAMINO A: PDF DIGITAL (Prioridad Nativa) ---
        if es_digital:
            print("PDF Digital detectado. Usando extracción Nativa (Alta Precisión)...")
            tables_found = self._extract_native_fallback()
            
            # Si el método nativo falló (ej: tablas sin bordes), intentamos el salvavidas de IA
            if not tables_found and self.usar_ia:
                print("Nativo no encontró tablas. Intentando con IA (GMFT)...")
                tables_found = self._extract_with_ai()

        # --- CAMINO B: PDF ESCANEADO (Prioridad IA/OCR) ---
        else:
            print("PDF Escaneado/Imagen detectado.")
            if self.usar_ia:
                print("Intentando IA (GMFT)...")
                tables_found = self._extract_with_ai()
            
            if not tables_found and self.has_ocr:
                print("IA falló. Usando Fuerza Bruta (OCR)...")
                tables_found = self._extract_with_vision_fallback()

        # --- GUARDADO ---
        if tables_found:
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                for sheet_name, df in tables_found:
                    counter = 1
                    base = sheet_name
                    # Evitar duplicados de nombre de hoja
                    while sheet_name in writer.book.sheetnames:
                        sheet_name = f"{base}_{counter}"
                        counter += 1
                    df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
            print(f"Excel generado con {len(tables_found)} tablas.")
            return excel_path
        else:
            print("No se pudo extraer ninguna tabla.")
            return None

    def _extract_with_ai(self):
        """Extrae usando GMFT"""
        found = []
        try:
            doc = PyPDFium2Document(self.file_path)
            for table in doc.tables():
                try:
                    df = self.formatter.extract(table)
                    if not df.empty and len(df) > 1:
                        df = df.map(lambda x: self._clean_text(str(x)) if x is not None else x)
                        found.append((f"P{table.page.page_number+1}_IA", df))
                except: pass
            doc.close()
        except Exception as e: 
            print(f"Error usando libreria GMFT: {e}")
        return found

    def _extract_native_fallback(self):
            tables = []
            try:
                with pdfplumber.open(self.file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        
                        # --- INTENTO 1: ESTRATEGIA DE LÍNEAS (Para tablas bien dibujadas) ---
                        extracted = page.extract_tables(table_settings={
                            "vertical_strategy": "lines", 
                            "horizontal_strategy": "lines",
                            "snap_tolerance": 3
                        })

                        # Validamos: Si no encontró nada O si encontró tablas pero de 1 sola columna
                        usar_estrategia_texto = False
                        if not extracted:
                            usar_estrategia_texto = True
                        else:
                            # Si la tabla detectada tiene menos de 2 columnas, probablemente es un falso positivo
                            if len(extracted[0][0]) < 2:
                                usar_estrategia_texto = True

                        # --- INTENTO 2: ESTRATEGIA DE TEXTO (Para tablas sin bordes) ---
                        if usar_estrategia_texto:
                            print(f" Pág {i+1}: Líneas no detectadas, buscando por espacios en blanco...")
                            extracted = page.extract_tables(table_settings={
                                "vertical_strategy": "text", 
                                "horizontal_strategy": "text",
                                "snap_tolerance": 3,
                                "intersection_x_tolerance": 15 # Tolerancia para alinear columnas chuecas
                            })

                        # Procesar lo que haya encontrado (sea por líneas o por texto)
                        for idx, t in enumerate(extracted):
                            if t and len(t) > 1: # Debe tener más de 1 fila para ser útil
                                
                                # Limpieza de Nulos
                                headers = t[0]
                                data = t[1:]
                                
                                # Crear nombres de columnas seguros
                                safe_headers = []
                                for k, col in enumerate(headers):
                                    col_str = str(col).strip() if col else ""
                                    # Si la columna no tiene nombre, le ponemos Col_X
                                    if not col_str:
                                        col_str = f"Col_{k+1}"
                                    # Si el nombre está repetido, le agregamos sufijo
                                    if col_str in safe_headers:
                                        col_str = f"{col_str}_{k+1}"
                                    safe_headers.append(col_str)

                                df = pd.DataFrame(data, columns=safe_headers)
                                
                                # Filtro final: Si el dataframe está casi vacío, lo ignoramos
                                if not df.dropna(how='all').empty:
                                    tables.append((f"P{i+1}_Nativa_{idx+1}", df))
                                    
            except Exception as e:
                print(f"   ⚠️ Error en extracción nativa: {e}")
                
            return tables


    def _extract_with_vision_fallback(self):
        found_tables = []
        try:
            with fitz.open(self.file_path) as doc:
                for i in range(len(doc)):
                    page = doc[i]
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))

                    data = pytesseract.image_to_data(image, lang='spa', output_type=pytesseract.Output.DATAFRAME)
                    data = data[data.conf > 30] 
                    data = data[data.text.notna()]
                    data['text'] = data['text'].astype(str).str.strip()
                    data = data[data['text'] != '']

                    if not data.empty:
                        data['row_group'] = (data['top'] // 15)
                        rows = []
                        for _, line in data.groupby('row_group'):
                            sorted_line = line.sort_values('left')
                            rows.append(sorted_line['text'].tolist())

                        if len(rows) > 1:
                            max_cols = max(len(r) for r in rows)
                            normalized_rows = [r + [''] * (max_cols - len(r)) for r in rows]
                            df = pd.DataFrame(normalized_rows)
                            found_tables.append((f"P{i+1}_OCR", df))
        except: pass
        return found_tables

    def extract_text_doc_smart(self):
        print("Generando documento Word...")
        docx_path = os.path.join(self.output_dir, f"{self.filename}_edit.docx")
        
        es_nativo = self._has_text_content()

        if es_nativo:
            print("PDF Digital detectado. Usando conversión de alta fidelidad.")
            try:
                cv = Converter(self.file_path)
                cv.convert(docx_path, start=0, end=None)
                cv.close()
                return docx_path
            except Exception as e:
                print(f"Falló pdf2docx: {e}. Intentando OCR...")
        
        if self.has_ocr:
            print("PDF Escaneado detectado. Usando OCR")
            doc = Document()
            doc.add_heading(f'Texto Extraído (OCR) - {self.filename}', 0)
            try:
                with fitz.open(self.file_path) as pdf:
                    for i, page in enumerate(pdf):
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        image = Image.open(io.BytesIO(pix.tobytes("png")))
                        texto = pytesseract.image_to_string(image, lang='spa')
                        doc.add_heading(f'Página {i+1}', level=1)
                        doc.add_paragraph(texto)
                        if i < len(pdf) - 1: doc.add_page_break()
                doc.save(docx_path)
                return docx_path
            except: pass
        return None

    def extract_images(self):
        print("Extrayendo imágenes...")
        img_count = 0
        images_dir = os.path.join(self.output_dir, "imagenes")
        os.makedirs(images_dir, exist_ok=True)
        with fitz.open(self.file_path) as doc:
            for i in range(len(doc)):
                for img_idx, img in enumerate(doc[i].get_images(full=True)):
                    try:
                        base = doc.extract_image(img[0])
                        with open(os.path.join(images_dir, f"P{i+1}_{img_idx}.{base['ext']}"), "wb") as f:
                            f.write(base["image"])
                        img_count += 1
                    except: pass
        return img_count